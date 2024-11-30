import os
import pandas as pd
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from statsmodels.tsa.api import VAR
from statsmodels.regression.linear_model import OLS  # Correct import for OLS
from docx import Document

# Load the dataset
df = pd.read_csv(r'C:\Users\CATE\PycharmProjects\data sudan.csv')

# Clean column names (remove spaces and replace with underscores)
df.columns = df.columns.str.strip().str.replace(' ', '_')


# Function to perform ADF test and return the result as a string
def adf_test(series, name):
    result = adfuller(series)
    output = (
        f"ADF Test for {name}:\n"
        f"ADF Statistic: {result[0]}\n"
        f"p-value: {result[1]}\n"
        f"Critical Values: {result[4]}\n"
        f"The series is likely stationary (reject H0)." if result[1] < 0.05 else
        f"The series is likely non-stationary (fail to reject H0).\n\n"
    )
    return result[1] < 0.05, output  # Return if stationary or not along with the output string


# Function for Engle-Granger Two-Step Cointegration Test
def engle_granger_cointegration_test(df):
    # Step 1: Perform OLS regression (GDP on Per_Capita)
    Y = df['GDP']
    X = df['Per_Capita']
    X = sm.add_constant(X)  # Add a constant to the independent variable for the regression
    model = OLS(Y, X).fit()  # Fit the OLS model

    # Step 2: Get residuals and perform ADF test on the residuals
    residuals = model.resid
    adf_stat, p_value, _, _, _, critical_values = adfuller(residuals)

    output = (
        f"Engle-Granger Cointegration Test:\n"
        f"ADF Statistic on Residuals: {adf_stat}\n"
        f"p-value: {p_value}\n"
        f"Critical Values: {critical_values}\n"
        f"The residuals are likely stationary (reject H0), indicating cointegration."
        if p_value < 0.05 else
        f"The residuals are likely non-stationary (fail to reject H0), indicating no cointegration.\n"
    )
    return output


# Function to perform VAR (Vector Autoregression)
def var_model(df):
    model = VAR(df)
    results = model.fit(maxlags=5)  # Correcting to use maxlags instead of lags
    output = (
        f"VAR Model Results (maxlags=5):\n"
        f"Selected Lags: {results.k_ar}\n"
        f"Summary: {results.summary()}\n"
    )
    return output


# Perform ADF test for 'GDP' and 'Per Capita'
gdp_stationary, output_gdp = adf_test(df['GDP'], 'GDP')
per_capita_stationary, output_per_capita = adf_test(df['Per_Capita'], 'Per Capita')

# Create a Word document to save the output
doc = Document()
doc.add_heading('ADF Test, Cointegration, and VAR Results', 0)

# Add ADF test results to the document
doc.add_paragraph(output_gdp)
doc.add_paragraph(output_per_capita)

# Perform Cointegration Test (Engle-Granger Test)
cointegration_results = engle_granger_cointegration_test(df[['GDP', 'Per_Capita']])
doc.add_paragraph(cointegration_results)

# Perform VAR Model (since the series are stationary)
var_results = var_model(df[['GDP', 'Per_Capita']])
doc.add_paragraph(var_results)

# Save the document to Desktop (adjusted for Windows Desktop path)
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop", "ADF_Cointegration_VAR_Results.docx")
doc.save(desktop_path)

print(f"Results saved to: {desktop_path}")






